import requests
import json
import sqlite3
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    MessageHandler,
    CallbackQueryHandler,
    ContextTypes,
    filters
)
from docx import Document
import os

# --- Настройки ---
TELEGRAM_TOKEN = os.environ.get("TELEGRAM_TOKEN")  # Ваш Telegram Bot Token
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY")  # OpenRouter API Key
DB_FILE = "books.db"

# --- Инициализация базы ---
conn = sqlite3.connect(DB_FILE)
cursor = conn.cursor()
cursor.execute("""
CREATE TABLE IF NOT EXISTS books (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    chat_id INTEGER,
    book_title TEXT,
    content TEXT
)
""")
conn.commit()

# --- Функция генерации текста ---
def generate_text(prompt, history=""):
    """Генерация текста через Kimi K2 на OpenRouter с контекстом книги"""
    full_prompt = (
        "Ты — «Создатель книг». Пиши на русском языке длинный, увлекательный текст по следующему запросу:\n"
        f"{prompt}\n\n"
        "Контекст предыдущих частей книги (если есть):\n"
        f"{history}\n\n"
        "Требования:\n"
        "1. Пиши литературно, с диалогами и описанием мира.\n"
        "2. Сохраняй стиль и персонажей.\n"
        "3. Не меняй жанр без запроса.\n"
        "4. Если продолжение книги, учитывай весь предыдущий текст.\n"
        "5. Создавай главы с заголовками при необходимости.\n"
        "6. Не повторяй предложения.\n"
    )

    response = requests.post(
        url="https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
        },
        data=json.dumps({
            "model": "moonshotai/kimi-k2:free",
            "messages": [
                {"role": "user", "content": full_prompt}
            ]
        })
    )
    result = response.json()
    try:
        return result['choices'][0]['message']['content']
    except KeyError:
        return "Ошибка генерации текста."

# --- Команды бота ---
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "Привет! Я — «Создатель книг».\n"
        "Команды:\n"
        "/newbook - создать новую книгу\n"
        "/mybooks - показать мои книги\n"
    )

async def new_book(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Введите название новой книги:")
    return

async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.message.chat_id
    text = update.message.text.strip()

    # Проверяем, создаем ли мы новую книгу
    if context.user_data.get("creating_book"):
        title = text
        # Добавляем книгу в базу
        cursor.execute("INSERT INTO books (chat_id, book_title, content) VALUES (?, ?, ?)", (chat_id, title, ""))
        conn.commit()
        context.user_data["creating_book"] = False
        context.user_data["current_book"] = cursor.lastrowid
        await update.message.reply_text(f"Книга '{title}' создана! Введите начало истории или тему книги:")
        return

    # Определяем текущую книгу
    book_id = context.user_data.get("current_book")
    if not book_id:
        await update.message.reply_text("Сначала создайте книгу через /newbook")
        return

    # Получаем предыдущий контент книги
    cursor.execute("SELECT content FROM books WHERE id=?", (book_id,))
    row = cursor.fetchone()
    history = row[0] if row else ""

    # Генерируем продолжение книги
    continuation = generate_text(text, history)

    # Сохраняем обновленный контент
    updated_content = history + "\n" + text + "\n" + continuation + "\n"
    cursor.execute("UPDATE books SET content=? WHERE id=?", (updated_content, book_id))
    conn.commit()

    # Кнопки для продолжения или сохранения
    keyboard = [
        [InlineKeyboardButton("Продолжить книгу", callback_data=f"continue_{book_id}")],
        [InlineKeyboardButton("Сохранить в Word", callback_data=f"save_{book_id}")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(continuation, reply_markup=reply_markup)

# --- Обработка кнопок ---
async def button(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    data = query.data
    chat_id = query.message.chat_id

    if data.startswith("continue_"):
        book_id = int(data.split("_")[1])
        context.user_data["current_book"] = book_id
        await query.edit_message_text("Введите текст для продолжения книги:")
    elif data.startswith("save_"):
        book_id = int(data.split("_")[1])
        cursor.execute("SELECT book_title, content FROM books WHERE id=?", (book_id,))
        row = cursor.fetchone()
        if not row:
            await query.edit_message_text("Книга не найдена.")
            return
        title, content = row
        # Создаем Word файл
        doc = Document()
        doc.add_paragraph(content)
        file_name = f"{title}.docx"
        doc.save(file_name)
        with open(file_name, "rb") as f:
            await query.message.reply_document(f, filename=file_name)

# --- Команда показать книги ---
async def my_books(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.message.chat_id
    cursor.execute("SELECT id, book_title FROM books WHERE chat_id=?", (chat_id,))
    rows = cursor.fetchall()
    if not rows:
        await update.message.reply_text("У вас нет созданных книг.")
        return
    text = "Ваши книги:\n"
    keyboard = []
    for book_id, title in rows:
        text += f"- {title}\n"
        keyboard.append([InlineKeyboardButton(title, callback_data=f"continue_{book_id}")])
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(text, reply_markup=reply_markup)

# --- Запуск бота ---
async def main():
    app = ApplicationBuilder().token(TELEGRAM_TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("newbook", new_book, block=False))
    app.add_handler(CommandHandler("mybooks", my_books))
    app.add_handler(MessageHandler(filters.TEXT & (~filters.COMMAND), handle_text))
    app.add_handler(CallbackQueryHandler(button))

    print("Бот запущен...")
    await app.run_polling()

if __name__ == "__main__":
    import asyncio
    asyncio.run(main())    return response.choices[0].message

def create_word_file(text, filename="book.docx"):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)
    return filename

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.message.from_user.id
    prompt = update.message.text

    await update.message.reply_text("Генерирую текст, это может занять минуту...")

    # Получаем историю пользователя или создаём новую
    history = user_histories.get(user_id, [])
    history.append({"role": "user", "content": prompt})

    # Генерация текста
    response_msg = generate_text(history)

    # Добавляем ответ в историю
    history.append({
        "role": "assistant",
        "content": response_msg.content,
        "reasoning_details": getattr(response_msg, "reasoning_details", None)
    })

    # Сохраняем обновлённую историю
    user_histories[user_id] = history

    # Создаём Word файл
    filename = create_word_file(response_msg.content)

    # Отправляем файл пользователю
    with open(filename, "rb") as f:
        await update.message.reply_document(InputFile(f, filename=filename))

# -----------------------
# Запуск бота
# -----------------------
if __name__ == "__main__":
    app = ApplicationBuilder().token(TELEGRAM_TOKEN).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.TEXT & (~filters.COMMAND), handle_message))
    print("Бот запущен...")
    app.run_polling()
